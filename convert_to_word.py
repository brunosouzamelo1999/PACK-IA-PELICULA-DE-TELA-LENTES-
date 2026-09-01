import sys
import os
import re
import subprocess

# Instala python-docx se necessário
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx nao encontrado. Instalando...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Define a cor de fundo de uma célula da tabela"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Define margens internas de uma celula (padding)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def format_bold_text(paragraph, text, base_font_name="Calibri", base_size=11, color_rgb=(51,51,51)):
    """Processa negritos markdown **texto** no paragrafo"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean_text = part[2:-2]
            run = paragraph.add_run(clean_text)
            run.bold = True
        else:
            run = paragraph.add_run(part)
        
        run.font.name = base_font_name
        run.font.size = Pt(base_size)
        run.font.color.rgb = RGBColor(*color_rgb)

def md_to_docx(md_path, docx_path):
    print(f"Lendo {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilos de cores
    PRIMARY_COLOR = RGBColor(16, 124, 65) # Verde MS Word
    TEXT_COLOR = RGBColor(51, 51, 51)     # Charcoal escuro
    GRAY_COLOR = RGBColor(120, 120, 120)
    
    in_table = False
    table_rows = []
    
    in_code_block = False
    code_block_text = []

    for line_idx, line in enumerate(lines):
        clean_line = line.strip()
        
        # Ignorar delimitadores de diagramas mermaid no word
        if clean_line.startswith('```mermaid'):
            in_code_block = True
            code_block_text = []
            continue
        
        if in_code_block:
            if clean_line.startswith('```'):
                in_code_block = False
                # Pula diagramas ou blocos de código extensos do word
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("[Fluxograma/Diagrama omitido no Word - Veja no manual markdown]")
                run.italic = True
                run.font.color.rgb = GRAY_COLOR
                run.font.size = Pt(9.5)
                continue
            code_block_text.append(line)
            continue
            
        # Pular blocos de prompt (fazer um tratamento especial de caixa de dica)
        if clean_line.startswith('```text') or clean_line.startswith('```'):
            if clean_line.startswith('```text') or clean_line.startswith('```'):
                if not in_table:
                    # Inicia ou encerra bloco de texto/código
                    pass
            continue
            
        # Identificar tabelas markdown
        if clean_line.startswith('|'):
            # Se for divisor de tabela (| :--- | :--- |), ignorar
            if re.match(r'^\|[\s:-|]+$', clean_line):
                continue
            
            in_table = True
            # Quebrar as colunas e limpar
            cols = [col.strip() for col in clean_line.split('|')[1:-1]]
            table_rows.append(cols)
            continue
        else:
            if in_table:
                # Gerar a tabela acumulada no docx
                if table_rows:
                    num_cols = len(table_rows[0])
                    doc_table = doc.add_table(rows=0, cols=num_cols)
                    doc_table.autofit = True
                    
                    # Adicionar linhas
                    for idx, row_cells in enumerate(table_rows):
                        row = doc_table.add_row()
                        for col_idx, text in enumerate(row_cells):
                            cell = row.cells[col_idx]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.space_before = Pt(2)
                            
                            # Cabecalho da tabela
                            if idx == 0:
                                set_cell_background(cell, "107C41") # Verde Word
                                set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
                                format_bold_text(p, f"**{text}**", base_size=10, color_rgb=(255,255,255))
                            else:
                                # Linhas alternadas zebra shading
                                if idx % 2 == 0:
                                    set_cell_background(cell, "F2F9F4")
                                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                                format_bold_text(p, text, base_size=9.5)
                    
                    doc.add_paragraph() # Espaco pos-tabela
                table_rows = []
                in_table = False

        # Títulos
        if clean_line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_line[2:])
            run.bold = True
            run.font.size = Pt(20)
            run.font.name = "Arial"
            run.font.color.rgb = PRIMARY_COLOR
            continue
            
        elif clean_line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(clean_line[3:])
            run.bold = True
            run.font.size = Pt(15)
            run.font.name = "Arial"
            run.font.color.rgb = PRIMARY_COLOR
            continue
            
        elif clean_line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(clean_line[4:])
            run.bold = True
            run.font.size = Pt(12.5)
            run.font.name = "Arial"
            run.font.color.rgb = PRIMARY_COLOR
            continue

        # Alertas e Dicas (> [!TIP], > [!IMPORTANT], > [!WARNING])
        if clean_line.startswith('>'):
            alert_text = clean_line[1:].strip()
            if alert_text.startswith('[!'):
                # Cria caixa de alerta como uma tabela de celula unica com bordas
                tbl = doc.add_table(rows=1, cols=1)
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "F4F4F4")
                set_cell_margins(cell, top=150, bottom=150, left=200, right=200)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                
                # Identifica tipo de alerta
                type_match = re.match(r'^\[!(TIP|IMPORTANT|WARNING|NOTE|CAUTION)\]', alert_text)
                if type_match:
                    alert_type = type_match.group(1)
                    clean_alert_text = alert_text[len(type_match.group(0)):].strip()
                    prefix = f"{alert_type}: "
                    run_pref = p.add_run(prefix)
                    run_pref.bold = True
                    run_pref.font.name = "Calibri"
                    run_pref.font.size = Pt(10)
                    if alert_type == "WARNING" or alert_type == "CAUTION":
                        run_pref.font.color.rgb = RGBColor(180, 50, 50)
                    else:
                        run_pref.font.color.rgb = PRIMARY_COLOR
                    
                    format_bold_text(p, clean_alert_text, base_size=10, color_rgb=(80,80,80))
                else:
                    format_bold_text(p, alert_text, base_size=10, color_rgb=(80,80,80))
                
                doc.add_paragraph()
                continue
            else:
                # É uma linha de continuação do bloco de citação
                # Adiciona ao último parágrafo da última tabela (se existir)
                if doc.tables:
                    last_tbl = doc.tables[-1]
                    cell = last_tbl.cell(0, 0)
                    p = cell.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    format_bold_text(p, alert_text, base_size=10, color_rgb=(80,80,80))
                continue

        # Listas com marcadores (bullet points)
        if clean_line.startswith('* ') or clean_line.startswith('- '):
            item_text = clean_line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            format_bold_text(p, item_text)
            continue
            
        # Listas numeradas
        num_match = re.match(r'^(\d+)\.\s(.*)$', clean_line)
        if num_match:
            num = num_match.group(1)
            item_text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            format_bold_text(p, item_text)
            continue

        # Parágrafo normal
        if clean_line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            format_bold_text(p, clean_line)
        else:
            # Linha vazia vira espaçamento
            pass

    print(f"Salvando {docx_path}...")
    doc.save(docx_path)
    print("Conversao concluida!")

if __name__ == "__main__":
    workspace = r"C:\Users\bruno\.gemini\antigravity-ide\scratch\workana_packaging_ai"
    
    # 1. Converter Manual Operacional
    md_file = os.path.join(workspace, "manual_operacional.md")
    docx_file = os.path.join(workspace, "manual_operacional.docx")
    md_to_docx(md_file, docx_file)
    
    # 2. Converter Plano de Implementacao
    brain_dir = r"C:\Users\bruno\.gemini\antigravity-ide\brain\cd626bd4-4389-4f71-b18f-a455298b060f"
    plan_md = os.path.join(brain_dir, "implementation_plan.md")
    plan_docx = os.path.join(workspace, "plano_de_implementacao.docx")
    if os.path.exists(plan_md):
        md_to_docx(plan_md, plan_docx)

